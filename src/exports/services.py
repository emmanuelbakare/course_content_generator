"""Creation and execution boundaries for private course exports."""

import io
import re
from dataclasses import dataclass
from html import escape

from django.core.exceptions import ValidationError
from django.core.files.base import ContentFile
from django.db import transaction
from django.utils import timezone
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate

from courses.models import Course, CurriculumVersion

from .models import ExportFile, ExportJob


class ExportConfigurationError(ValidationError):
    """Raised when a course cannot safely be exported."""


@dataclass(frozen=True)
class CourseExportContent:
    title: str
    topic: str
    description: str
    sections: tuple


def enqueue_course_export(*, course_id, requested_by, export_format: str) -> ExportJob:
    """Queue an export of the sole approved curriculum snapshot for its owner."""
    course = Course.objects.get(pk=course_id, owner=requested_by)
    curriculum = course.curriculum_versions.filter(
        status=CurriculumVersion.Status.APPROVED
    ).first()
    if not curriculum:
        raise ExportConfigurationError('Approve a curriculum before exporting the course.')
    job = ExportJob(
        course=course,
        curriculum_version=curriculum,
        requested_by=requested_by,
        export_format=export_format,
    )
    job.full_clean()
    job.save()
    from .tasks import run_export_job

    run_export_job.delay(str(job.public_id))
    return job


def process_export_job(job_id) -> ExportJob:
    """Render and persist a user-private export without involving LLM services."""
    with transaction.atomic():
        job = ExportJob.objects.select_for_update().select_related(
            'course', 'curriculum_version', 'requested_by'
        ).get(public_id=job_id)
        if job.status == ExportJob.Status.CANCELLED:
            return job
        if job.status == ExportJob.Status.SUCCEEDED:
            return job
        job.status = ExportJob.Status.RUNNING
        job.started_at = timezone.now()
        job.error_message = ''
        job.save(update_fields=('status', 'started_at', 'error_message'))

    try:
        content = _load_export_content(job)
        filename, content_type, payload = _render_export(content, job.export_format)
        filename = f'{_slugify(job.course.title)}.{filename}'
        with transaction.atomic():
            job = ExportJob.objects.select_for_update().get(pk=job.pk)
            if job.status == ExportJob.Status.CANCELLED:
                return job
            export_file, _ = ExportFile.objects.get_or_create(
                job=job,
                defaults={'owner': job.requested_by, 'original_filename': filename, 'content_type': content_type},
            )
            export_file.owner = job.requested_by
            export_file.original_filename = filename
            export_file.content_type = content_type
            export_file.file.save(filename, ContentFile(payload), save=False)
            export_file.size_bytes = len(payload)
            export_file.full_clean()
            export_file.save()
            job.status = ExportJob.Status.SUCCEEDED
            job.progress_current = job.progress_total
            job.completed_at = timezone.now()
            job.save(update_fields=('status', 'progress_current', 'completed_at'))
            return job
    except Exception as exc:
        ExportJob.objects.filter(pk=job.pk).update(
            status=ExportJob.Status.FAILED, error_message=str(exc), completed_at=timezone.now()
        )
        raise


def _load_export_content(job: ExportJob) -> CourseExportContent:
    curriculum = CurriculumVersion.objects.prefetch_related(
        'sections__lessons__revisions'
    ).get(pk=job.curriculum_version_id)
    sections = []
    for section in curriculum.sections.all():
        lessons = []
        for lesson in section.lessons.all():
            latest_revision = lesson.revisions.first()
            lessons.append((lesson.title, latest_revision.content_markdown if latest_revision else lesson.outline))
        sections.append((section.title, section.summary, tuple(lessons)))
    return CourseExportContent(job.course.title, job.course.topic, curriculum.course_description, tuple(sections))


def _render_export(content: CourseExportContent, export_format: str):
    if export_format == ExportJob.Format.MARKDOWN:
        return 'md', 'text/markdown; charset=utf-8', _render_markdown(content).encode('utf-8')
    if export_format == ExportJob.Format.DOCX:
        return 'docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', _render_docx(content)
    if export_format == ExportJob.Format.PDF:
        return 'pdf', 'application/pdf', _render_pdf(content)
    raise ExportConfigurationError('Unsupported export format.')


def _render_markdown(content: CourseExportContent) -> str:
    lines = [f'# {content.title}', '', f'**Topic:** {content.topic}']
    if content.description:
        lines.extend(['', content.description])
    for section_title, summary, lessons in content.sections:
        lines.extend(['', f'## {section_title}'])
        if summary:
            lines.extend(['', summary])
        for lesson_title, lesson_content in lessons:
            lines.extend(['', f'### {lesson_title}', '', lesson_content or '_Content has not been written yet._'])
    return '\n'.join(lines).strip() + '\n'


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    normal = document.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ('Heading 1', 16, '2E74B5', 18, 10),
        ('Heading 2', 13, '2E74B5', 14, 7),
        ('Heading 3', 12, '1F4D78', 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    title = document.styles['Title']
    title.font.name = 'Calibri'
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor.from_string('0B2545')
    title.paragraph_format.space_after = Pt(12)


def _render_docx(content: CourseExportContent) -> bytes:
    document = Document()
    _configure_docx(document)  # compact_reference_guide token set
    document.add_paragraph(content.title, style='Title')
    document.add_paragraph(f'Topic: {content.topic}')
    if content.description:
        document.add_paragraph(content.description)
    for section_title, summary, lessons in content.sections:
        document.add_heading(section_title, level=1)
        if summary:
            document.add_paragraph(summary)
        for lesson_title, lesson_content in lessons:
            document.add_heading(lesson_title, level=2)
            _append_markdown_to_docx(document, lesson_content or 'Content has not been written yet.')
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _append_markdown_to_docx(document: Document, markdown_text: str) -> None:
    for line in markdown_text.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith('### '):
            document.add_heading(text[4:], level=3)
        elif text.startswith('## '):
            document.add_heading(text[3:], level=2)
        elif text.startswith('# '):
            document.add_heading(text[2:], level=1)
        elif text.startswith(('- ', '* ')):
            document.add_paragraph(_plain_markdown(text[2:]), style='List Bullet')
        else:
            document.add_paragraph(_plain_markdown(text))


def _render_pdf(content: CourseExportContent) -> bytes:
    output = io.BytesIO()
    styles = getSampleStyleSheet()
    body = ParagraphStyle('CourseBody', parent=styles['BodyText'], fontName='Helvetica', fontSize=10.5,
                          leading=14, spaceAfter=7, alignment=TA_LEFT)
    title = ParagraphStyle('CourseTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=22,
                           leading=27, textColor=HexColor('#0B2545'), spaceAfter=12)
    h1 = ParagraphStyle('CourseH1', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=16,
                        leading=20, textColor=HexColor('#2E74B5'), spaceBefore=16, spaceAfter=8)
    h2 = ParagraphStyle('CourseH2', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=13,
                        leading=16, textColor=HexColor('#2E74B5'), spaceBefore=12, spaceAfter=6)
    h3 = ParagraphStyle('CourseH3', parent=styles['Heading3'], fontName='Helvetica-Bold', fontSize=11.5,
                        leading=14, textColor=HexColor('#1F4D78'), spaceBefore=8, spaceAfter=4)
    story = [Paragraph(escape(content.title), title), Paragraph(f'<b>Topic:</b> {escape(content.topic)}', body)]
    if content.description:
        story.append(Paragraph(escape(content.description).replace('\n', '<br/>'), body))
    for section_title, summary, lessons in content.sections:
        story.append(Paragraph(escape(section_title), h1))
        if summary:
            story.append(Paragraph(escape(summary), body))
        for lesson_title, lesson_content in lessons:
            story.append(Paragraph(escape(lesson_title), h2))
            _append_markdown_to_pdf(story, lesson_content or 'Content has not been written yet.', body, h1, h2, h3)
    SimpleDocTemplate(output, pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch).build(story)
    return output.getvalue()


def _append_markdown_to_pdf(story, markdown_text, body, h1, h2, h3):
    for line in markdown_text.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith('### '):
            story.append(Paragraph(escape(text[4:]), h3))
        elif text.startswith('## '):
            story.append(Paragraph(escape(text[3:]), h2))
        elif text.startswith('# '):
            story.append(Paragraph(escape(text[2:]), h1))
        else:
            prefix = '&bull; ' if text.startswith(('- ', '* ')) else ''
            story.append(Paragraph(prefix + escape(_plain_markdown(text[2:] if prefix else text)), body))


def _plain_markdown(value: str) -> str:
    return re.sub(r'(?<!\\)[*_`]+', '', value)


def _slugify(value: str) -> str:
    slug = re.sub(r'[^a-z0-9]+', '-', value.lower()).strip('-')
    return slug or 'course-export'
